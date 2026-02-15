"""
Document converter — extracts text from various file formats.

Handles PDF, Office documents, HTML, and plain text.  Uses pure-Python
libraries with optional Tika fallback for complex formats.
"""

import io
import logging
import mimetypes
import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional

logger = logging.getLogger("sf.enrichment.converter")


@dataclass
class ConversionResult:
    """Result of document text extraction."""

    text: str = ""
    metadata: Dict[str, str] = field(default_factory=dict)
    pages: int = 0
    format_detected: str = "unknown"
    language: str = "unknown"
    errors: List[str] = field(default_factory=list)
    char_count: int = 0

    @property
    def is_success(self) -> bool:
        return len(self.text) > 0


class DocumentConverter:
    """
    Converts various document formats to plain text.

    Tries pure-Python extractors first, falls back to Apache Tika
    if configured (via SF_TIKA_ENDPOINT).
    """

    def __init__(self):
        self.tika_endpoint = os.environ.get("SF_TIKA_ENDPOINT", "")

    def convert(self, content: bytes, filename: str, content_type: str = "") -> ConversionResult:
        """
        Convert document content to plain text.

        Args:
            content: Raw file bytes
            filename: Original filename
            content_type: MIME type (auto-detected if empty)

        Returns:
            ConversionResult with extracted text and metadata
        """
        if not content_type:
            content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

        result = ConversionResult(format_detected=content_type)

        try:
            if content_type == "application/pdf":
                result = self._convert_pdf(content, result)
            elif content_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ):
                result = self._convert_docx(content, result)
            elif content_type in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ):
                result = self._convert_xlsx(content, result)
            elif content_type in ("text/html", "application/xhtml+xml"):
                result = self._convert_html(content, result)
            elif content_type.startswith("text/") or content_type in (
                "application/json",
                "application/xml",
                "application/javascript",
            ):
                result = self._convert_text(content, result)
            elif content_type == "application/rtf":
                result = self._convert_rtf(content, result)
            elif self.tika_endpoint:
                result = self._convert_tika(content, content_type, result)
            else:
                # Last resort: try as text
                result = self._convert_text(content, result)

        except Exception as exc:
            result.errors.append(f"Conversion error: {exc}")
            logger.warning("Conversion failed for %s: %s", filename, exc)

        result.char_count = len(result.text)
        return result

    def _convert_pdf(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            result.text = "\n\n".join(pages)
            result.pages = len(reader.pages)
            result.format_detected = "application/pdf"

            # Extract metadata
            meta = reader.metadata
            if meta:
                for key in ("title", "author", "creator", "producer", "subject"):
                    val = getattr(meta, key, None)
                    if val:
                        result.metadata[key] = str(val)

        except ImportError:
            result.errors.append("pypdf not installed — PDF extraction unavailable")
            logger.warning("pypdf not available for PDF conversion")

        return result

    def _convert_docx(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Extract text from DOCX."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            result.text = "\n".join(paragraphs)
            result.format_detected = "application/docx"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        result.text += "\n" + " | ".join(cells)

            # Core properties
            props = doc.core_properties
            if props.title:
                result.metadata["title"] = props.title
            if props.author:
                result.metadata["author"] = props.author

        except ImportError:
            result.errors.append("python-docx not installed")

        return result

    def _convert_xlsx(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Extract text from Excel spreadsheets."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append(" | ".join(cells))

            result.text = "\n".join(lines)
            result.format_detected = "application/xlsx"
            result.pages = len(wb.sheetnames)

        except ImportError:
            result.errors.append("openpyxl not installed")

        return result

    def _convert_html(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Extract text from HTML."""
        try:
            from html.parser import HTMLParser

            class _TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_parts = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style", "noscript"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style", "noscript"):
                        self._skip = False
                    if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "li", "tr"):
                        self.text_parts.append("\n")

                def handle_data(self, data):
                    if not self._skip:
                        text = data.strip()
                        if text:
                            self.text_parts.append(text)

            decoded = content.decode("utf-8", errors="replace")
            parser = _TextExtractor()
            parser.feed(decoded)
            result.text = " ".join(parser.text_parts)
            result.format_detected = "text/html"

        except Exception as exc:
            result.errors.append(f"HTML parse error: {exc}")

        return result

    def _convert_text(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Handle plain text content."""
        for encoding in ("utf-8", "latin-1", "ascii"):
            try:
                result.text = content.decode(encoding)
                result.format_detected = "text/plain"
                break
            except (UnicodeDecodeError, ValueError):
                continue

        if not result.text:
            result.text = content.decode("utf-8", errors="replace")
            result.errors.append("Decoded with replacement characters")

        return result

    def _convert_rtf(self, content: bytes, result: ConversionResult) -> ConversionResult:
        """Extract text from RTF (basic)."""
        try:
            from striprtf.striprtf import rtf_to_text

            decoded = content.decode("utf-8", errors="replace")
            result.text = rtf_to_text(decoded)
            result.format_detected = "application/rtf"
        except ImportError:
            # Fallback: strip RTF control words
            decoded = content.decode("utf-8", errors="replace")
            result.text = re.sub(r'[{}\\\*]|\\[a-z]+\d* ?', '', decoded)
            result.format_detected = "application/rtf"
            result.errors.append("striprtf not installed — basic RTF stripping used")

        return result

    def _convert_tika(
        self, content: bytes, content_type: str, result: ConversionResult
    ) -> ConversionResult:
        """Fallback: use Apache Tika for conversion."""
        try:
            import requests

            resp = requests.put(
                f"{self.tika_endpoint}/tika",
                data=content,
                headers={"Content-Type": content_type, "Accept": "text/plain"},
                timeout=60,
            )
            resp.raise_for_status()
            result.text = resp.text
            result.format_detected = content_type

        except ImportError:
            result.errors.append("requests library not available for Tika fallback")
        except Exception as exc:
            result.errors.append(f"Tika conversion failed: {exc}")

        return result
